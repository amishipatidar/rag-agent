"""
Script to generate sample test files for Phase 2 tests.
Run this once: python tests/create_samples.py
"""
import os
import sys

# Add project root to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

SAMPLES_DIR = os.path.join(os.path.dirname(__file__), "samples")
os.makedirs(SAMPLES_DIR, exist_ok=True)


def create_sample_docx():
    """Create a sample .docx file."""
    from docx import Document
    doc = Document()
    doc.add_heading("RAGgent Test Document", level=1)
    doc.add_paragraph(
        "This is a sample document for testing the RAGgent document parser. "
        "It contains multiple paragraphs with different topics."
    )
    doc.add_paragraph(
        "Machine learning is a subset of artificial intelligence that enables "
        "systems to learn and improve from experience without being explicitly programmed."
    )
    doc.add_paragraph(
        "Natural language processing allows computers to understand, interpret, "
        "and generate human language in a valuable way."
    )
    path = os.path.join(SAMPLES_DIR, "sample.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_sample_xlsx():
    """Create a sample .xlsx file."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Products"

    ws.append(["Product", "Category", "Price"])
    ws.append(["Laptop", "Electronics", 999.99])
    ws.append(["Headphones", "Electronics", 149.99])
    ws.append(["Notebook", "Stationery", 4.99])

    path = os.path.join(SAMPLES_DIR, "sample.xlsx")
    wb.save(path)
    print(f"Created: {path}")


def create_sample_pdf():
    """Create a simple sample .pdf file using reportlab if available, else a minimal PDF."""
    path = os.path.join(SAMPLES_DIR, "sample.pdf")
    try:
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(path)
        c.drawString(72, 750, "RAGgent PDF Test Document")
        c.drawString(72, 720, "This is a test PDF for the document parser.")
        c.drawString(72, 690, "Vector databases store embeddings for semantic search.")
        c.save()
    except ImportError:
        # Fallback: create a minimal valid PDF manually
        pdf_content = (
            "%PDF-1.4\n"
            "1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
            "2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
            "3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            "/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
            "4 0 obj\n<< /Length 44 >>\nstream\nBT /F1 12 Tf 72 750 Td (RAGgent PDF Test) Tj ET\n"
            "endstream\nendobj\n"
            "5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
            "xref\n0 6\n"
            "0000000000 65535 f \n"
            "0000000009 00000 n \n"
            "0000000058 00000 n \n"
            "0000000115 00000 n \n"
            "0000000282 00000 n \n"
            "0000000380 00000 n \n"
            "trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n456\n%%EOF\n"
        )
        with open(path, "w") as f:
            f.write(pdf_content)
    print(f"Created: {path}")


if __name__ == "__main__":
    create_sample_docx()
    create_sample_xlsx()
    create_sample_pdf()
    print("\nAll sample files created successfully!")
